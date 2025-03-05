import streamlit as st
import pandas as pd
import camelot
import tempfile
import os
import re
from PyPDF2 import PdfReader
import base64
import pdfplumber
from datetime import datetime
from io import BytesIO

# Para o fuzzy matching
from fuzzywuzzy import process, fuzz

# Para gerar PDFs
from fpdf import FPDF

# Para gerar DOCX
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

# -----------------------------------------------------------------
# CONFIGURAÇÕES DE STREAMLIT
# -----------------------------------------------------------------
st.set_page_config(
    page_title="Analisador de Contracheques - SEAD / AMAZONPREV",
    layout="centered"
)

LOGO_PATH = "MP.png"  # Ajuste se necessário
GLOSSARY_PATH = "Rubricas.txt"  # Ajuste se necessário

# Variável para armazenar valores importantes (fallback quando session_state não estiver disponível)
_fallback_state = {
    "df_contracheques": None,
    "df_descontos": None,
    "df_descontos_gloss": None,
    "df_descontos_gloss_sel": None,
    "nome_extraido": "",
    "nit_extraido": "",
    "valor_recebido": ""  # Para inserir o valor B = Valor Recebido
}


def get_state_value(key):
    """
    Recupera valor do st.session_state, se existir;
    caso contrário, busca em _fallback_state.
    """
    try:
        return st.session_state[key]
    except:
        return _fallback_state.get(key, None)


def set_state_value(key, value):
    """
    Armazena valor em st.session_state (se disponível) ou em _fallback_state.
    """
    try:
        st.session_state[key] = value
    except:
        _fallback_state[key] = value


# -----------------------------------------------------------------
# FUNÇÕES GERAIS
# -----------------------------------------------------------------
def get_image_base64(file_path):
    """Carrega imagem e retorna string base64 para exibir no Streamlit."""
    if not os.path.exists(file_path):
        return ""
    with open(file_path, "rb") as img_file:
        return base64.b64encode(img_file.read()).decode()


def carregar_glossario(path):
    """Carrega Rubricas de um arquivo texto."""
    try:
        with open(path, "r", encoding="utf-8") as f:
            return f.read().splitlines()
    except Exception as e:
        st.error(f"Erro ao carregar glossário: {e}")
        return []


def carregar_glossario_rubricas():
    """Carrega a lista de rubricas do GLOSSARY_PATH."""
    return carregar_glossario(GLOSSARY_PATH)


def sanitizar_para_arquivo(texto: str) -> str:
    """
    Remove caracteres indesejados para uso em nomes de arquivos.
    """
    texto = texto.strip().replace(" ", "_")
    return re.sub(r"[^\w\-_\.]", "", texto, flags=re.UNICODE)


def formatar_valor_brl(valor: str) -> str:
    """
    Converte valores do tipo "123,456.78" (ou variações) para padrão PT-BR "123.456,78".
    """
    try:
        flt = float(valor.replace(",", "").replace(".", "")) / 100
        return f"{flt:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    except:
        return valor


# -----------------------------------------------------------------
# EXTRAIR NOME E MATRÍCULA (via pdfplumber)
# -----------------------------------------------------------------
def extrair_nome_e_matricula(pdf_path):
    """
    Extrai Nome e Matrícula (MATRÍCULA-SEQ-DIG) do PDF usando pdfplumber.
    Exemplo de trecho que pode existir no PDF:
      NOME
      FULANO DE TAL
      MATRÍCULA-SEQ-DIG
      014.642-0 C
    Retorna: (nome, matrícula).
    Se não encontrar, retorna ("N/D", "N/D").
    """
    nome = "N/D"
    matricula = "N/D"
    with pdfplumber.open(pdf_path) as pdf:
        if len(pdf.pages) > 0:
            text = pdf.pages[0].extract_text() or ""
            lines = text.split("\n")

            for i, linha in enumerate(lines):
                if "NOME" in linha.upper():
                    if i + 1 < len(lines):
                        valor_nome = lines[i + 1].strip()
                        # Pega apenas parte textual sem números, se houver
                        match_nome = re.match(r"([^\d]+)", valor_nome)
                        if match_nome:
                            nome = match_nome.group(1).strip()
                if "MATRÍCULA-SEQ-DIG" in linha.upper():
                    if i + 1 < len(lines):
                        valor_matr = lines[i + 1].strip()
                        # Tenta casar algo como 014.642-0 C
                        matr_match = re.search(r"(\d{3}\.\d{3}-\d\s*[A-Z]*)", valor_matr)
                        if matr_match:
                            matricula = matr_match.group(1).strip()

    return nome or "N/D", matricula or "N/D"


# -----------------------------------------------------------------
# FUNÇÕES AUXILIARES CAMELOT
# -----------------------------------------------------------------
def extrair_data_da_pagina(pdf_path, page_number):
    """
    Utiliza PyPDF2 para extrair texto cru e procurar datas (MM/AAAA) em uma página específica.
    """
    with open(pdf_path, 'rb') as f:
        reader = PdfReader(f)
        if page_number - 1 < len(reader.pages):
            text = reader.pages[page_number - 1].extract_text() or ""
            match = re.search(r"\d{2}/\d{4}", text)
            if match:
                return match.group(0)
    return "N/D"


def _separar_linhas_multiplas(df: pd.DataFrame) -> pd.DataFrame:
    """
    Separa texto com quebras de linha (\n) em múltiplas linhas do DataFrame.
    """
    linhas_expandidas = []
    for _, row in df.iterrows():
        col_split = [str(row[col]).split('\n') for col in df.columns]
        max_splits = max(len(c) for c in col_split)
        for i in range(max_splits):
            nova_linha = {}
            for c, col_name in enumerate(df.columns):
                partes = col_split[c]
                nova_linha[col_name] = partes[i].strip() if i < len(partes) else ''
            linhas_expandidas.append(nova_linha)
    return pd.DataFrame(linhas_expandidas)


def processar_contracheques_camelot(pdf_path):
    """
    Lê tabelas do PDF com Camelot (flavor='stream') e retorna DataFrame
    com colunas: "COD", "Descrição", "TOTAL", "DATA".
    """
    try:
        tables = camelot.read_pdf(
            pdf_path,
            pages="all",
            flavor="stream",
            row_tol=15,
            strip_text=''
        )
        dados = pd.DataFrame(
            columns=["DATA", "Competência", "Descrição", "PVD", "COD", "BASE", "VALOR UNITÁRIO", "TOTAL"]
        )

        for table in tables:
            df = table.df
            # Verifica se existe a linha de cabeçalho com "DESCRIÇÃO"
            if "DESCRIÇÃO" in df.values:
                idx_cab = df[df.isin(["DESCRIÇÃO"]).any(axis=1)].index[0]
                df = df.iloc[idx_cab + 1:].reset_index(drop=True)
                if df.shape[1] >= 6:
                    df.columns = ["Descrição", "PVD", "COD", "BASE", "VALOR UNITÁRIO", "TOTAL"][:df.shape[1]]
                else:
                    continue

                # Separa as linhas com \n
                df = _separar_linhas_multiplas(df)

                # Marca de qual página veio a tabela (Competência = Página X)
                df["Competência"] = f"Página {table.page}"
                data_encontrada = extrair_data_da_pagina(pdf_path, table.page)
                df.insert(0, "DATA", data_encontrada)  # Insere DATA logo na 1ª posição
                dados = pd.concat([dados, df], ignore_index=True)

        if not dados.empty:
            # Reordena colunas para exibir final
            dados = dados[["COD", "Descrição", "TOTAL", "DATA"]]
        return dados
    except Exception as e:
        st.error(f"Erro ao processar o PDF: {str(e)}")
        return None


# -----------------------------------------------------------------
# PDF “Tabelas (SEAD / AMAZONPREV)”
# -----------------------------------------------------------------
class PDFRelatorioCamelot(FPDF):
    """
    Gera PDF em modo paisagem para exibir tabelas da SEAD / AMAZONPREV:
    Colunas: COD(30), DESCRIÇÃO(150), TOTAL(40), DATA(30).
    Possui cabeçalho e rodapé simples.
    """

    def __init__(self, titulo, nome_user, nit_user):
        super().__init__(orientation='L', unit='mm', format='A4')
        self.titulo = titulo
        self.nome_user = nome_user
        self.nit_user = nit_user
        # Ajustes de margem/página
        self.set_auto_page_break(auto=False, margin=15)
        self.set_left_margin(10)
        self.set_right_margin(10)
        self.set_top_margin(10)
        self.page_count = 0

    def header(self):
        self.page_count += 1
        self.set_font('Arial', 'B', 12)
        titulo_str = f"{self.titulo} - {self.nome_user} - {self.nit_user}"
        self.cell(0, 10, titulo_str, border=False, ln=True, align='C')
        self.ln(5)

        # Cabeçalho de colunas
        self._draw_header_line()

    def _draw_header_line(self):
        self.set_font("Arial", "B", 10)
        self.set_fill_color(200, 220, 255)
        header_cols = [
            ("COD", 30, "C"),
            ("Descrição", 150, "L"),
            ("TOTAL", 40, "R"),
            ("DATA", 30, "C")
        ]
        for col_name, col_w, _alig in header_cols:
            self.cell(col_w, 8, col_name, border=1, align='C', fill=True)
        self.ln()

    def footer(self):
        self.set_y(-15)
        self.set_font("Arial", "I", 8)
        self.cell(0, 10, f'Página {self.page_no()}', ln=False, align='C')

    def montar_tabela(self, df):
        """
        Escreve as linhas do DataFrame no PDF.
        """
        row_h = 7
        self.set_font("Arial", "", 9)

        for _, rowv in df.iterrows():
            if self.get_y() + row_h + 15 > self.h:  # Evitar estouro de página
                self.add_page()

            cod_s = str(rowv["COD"])
            desc_s = str(rowv["Descrição"])
            tot_s = str(rowv["TOTAL"])
            dat_s = str(rowv["DATA"])

            # Formata campo de total
            if tot_s.strip():
                tot_s = formatar_valor_brl(tot_s)

            self.cell(30, row_h, cod_s, border=1, align='C')
            self.cell(150, row_h, desc_s, border=1, align='L')
            self.cell(40, row_h, tot_s, border=1, align='R')
            self.cell(30, row_h, dat_s, border=1, align='C')
            self.ln(row_h)

    def gerar_pdf(self, df, out_path):
        self.add_page()
        self.montar_tabela(df)
        self.output(out_path)


def salvar_em_pdf_camelot(df, titulo_pdf, nome_user, nit_user) -> bytes:
    """
    Gera o PDF final (em bytes) das tabelas (SEAD / AMAZONPREV).
    """
    pdf = PDFRelatorioCamelot(titulo_pdf, nome_user, nit_user)
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_pdf:
        out_path = tmp_pdf.name
    pdf.gerar_pdf(df, out_path)
    with open(out_path, "rb") as f:
        data_bytes = f.read()
    os.remove(out_path)
    return data_bytes


# -----------------------------------------------------------------
# FUNÇÃO para inserir linhas de total A/B e cálculo de indébito
# -----------------------------------------------------------------
def inserir_totais_na_coluna(df, col_valor="DESCONTOS"):
    """
    Insere linhas ao final da coluna 'col_valor' com:
       - A = Valor Total (R$)
       - B = Valor Recebido - Autor (a)
       - Indébito (A-B)
       - Indébito em dobro (R$)

    O valor B é recuperado do estado ("valor_recebido").
    Se não fornecido ou inválido, considera 0.
    """

    def _to_float(x):
        try:
            return float(str(x).replace(',', '.').strip())
        except:
            return 0.0

    # Soma de A
    soma = df[col_valor].apply(_to_float).sum()
    if soma == 0:
        return df

    # Recupera B do estado
    valor_recebido_str = get_state_value("valor_recebido") or "0"
    try:
        valor_recebido_num = float(valor_recebido_str.replace(',', '.').strip())
    except:
        valor_recebido_num = 0.0

    # Indébito
    indebito = soma - valor_recebido_num
    indebito_dobro = 2 * indebito

    # Formata para exibir
    A_str = f"{soma:,.2f}"
    B_str = valor_recebido_str.strip()
    indebito_str = f"{indebito:,.2f}"
    indebito_dobro_str = f"{indebito_dobro:,.2f}"

    df_novo = df.copy()
    # Linha A
    df_novo = pd.concat([
        df_novo,
        pd.DataFrame({col_valor: [A_str], "DESCRIÇÃO": ["A = Valor Total (R$)"]})
    ], ignore_index=True)
    # Linha B
    df_novo = pd.concat([
        df_novo,
        pd.DataFrame({col_valor: [B_str], "DESCRIÇÃO": ["B = Valor Recebido - Autor (a)"]})
    ], ignore_index=True)
    # Linha Indébito
    df_novo = pd.concat([
        df_novo,
        pd.DataFrame({col_valor: [indebito_str], "DESCRIÇÃO": ["Indébito (A-B)"]})
    ], ignore_index=True)
    # Linha Indébito em dobro
    df_novo = pd.concat([
        df_novo,
        pd.DataFrame({col_valor: [indebito_dobro_str], "DESCRIÇÃO": ["Indébito em dobro (R$)"]})
    ], ignore_index=True)

    # Limpa demais colunas nas linhas especiais
    mask_especial = df_novo["DESCRIÇÃO"].isin([
        "A = Valor Total (R$)",
        "B = Valor Recebido - Autor (a)",
        "Indébito (A-B)",
        "Indébito em dobro (R$)"
    ])
    for c in df_novo.columns:
        if c not in ["DESCRIÇÃO", col_valor]:
            df_novo.loc[mask_especial, c] = ""

    return df_novo


# -----------------------------------------------------------------
# RELATÓRIO PDF “Descontos Finais”
# -----------------------------------------------------------------
class PDFFinais(FPDF):
    """
    PDF p/ "Descontos Finais", paisagem.
    Destaca linhas especiais:
      - A = Valor Total (R$)
      - B = Valor Recebido - Autor (a)
      - Indébito (A-B)
      - Indébito em dobro (R$)
    em vermelho (fonte 11, negrito).
    """

    def __init__(self, titulo):
        super().__init__(orientation='L', unit='mm', format='A4')
        self.titulo = titulo
        self.set_auto_page_break(auto=False, margin=15)
        self.set_left_margin(10)
        self.set_right_margin(10)
        self.set_top_margin(10)

    def header(self):
        self.set_font("Arial", "B", 12)
        self.cell(0, 10, self.titulo, border=False, ln=True, align='C')
        self.ln(5)
        self.set_font("Arial", "B", 10)
        self.set_fill_color(200, 220, 255)

    def footer(self):
        self.set_y(-15)
        self.set_font("Arial", "I", 8)
        self.cell(0, 10, f'Página {self.page_no()}', ln=False, align='C')

    def montar_cab(self, col_names, widths_map):
        row_h = 8
        self.set_font("Arial", "B", 10)
        self.set_fill_color(200, 220, 255)
        for col_h in col_names:
            w = widths_map.get(col_h, 40)
            self.cell(w, row_h, col_h, border=1, align='C', fill=True)
        self.ln()

    def montar_tabela(self, df):
        row_h = 7
        col_names = df.columns.tolist()
        widths_map = {
            "COD": 30,
            "DESCRIÇÃO": 120,
            "DESCONTOS": 40,
            "DATA": 30
        }
        # Desenha cabeçalho
        self.montar_cab(col_names, widths_map)
        self.set_font("Arial", "", 9)

        for _, rowv in df.iterrows():
            if self.get_y() + row_h + 15 > self.h:
                self.add_page()
                self.montar_cab(col_names, widths_map)

            desc = rowv["DESCRIÇÃO"]
            is_especial = desc in [
                "A = Valor Total (R$)",
                "B = Valor Recebido - Autor (a)",
                "Indébito (A-B)",
                "Indébito em dobro (R$)"
            ]

            if is_especial:
                self.set_font("Arial", "B", 11)
                self.set_text_color(255, 0, 0)
            else:
                self.set_font("Arial", "", 9)
                self.set_text_color(0, 0, 0)

            row_vals = []
            for col_h in col_names:
                val = str(rowv[col_h]) if col_h in rowv else ""
                if col_h == "DESCONTOS" and val.strip():
                    val = formatar_valor_brl(val)
                row_vals.append(val)

            # Imprime linha
            for col_h, valv in zip(col_names, row_vals):
                w = widths_map.get(col_h, 40)
                align = 'C'
                if col_h.upper() == "DESCRIÇÃO":
                    align = 'L'
                self.cell(w, row_h, valv, border=1, align=align)
            self.ln(row_h)

            if is_especial:
                self.set_font("Arial", "", 9)
                self.set_text_color(0, 0, 0)

    def gerar_pdf(self, df, out_path):
        self.add_page()
        self.montar_tabela(df)
        self.output(out_path)


def gerar_pdf_finais(df: pd.DataFrame, titulo: str) -> bytes:
    pdf = PDFFinais(titulo)
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmpf:
        out_path = tmpf.name
    pdf.gerar_pdf(df, out_path)
    with open(out_path, "rb") as f:
        data_bytes = f.read()
    os.remove(out_path)
    return data_bytes


# -----------------------------------------------------------------
# RELATÓRIO DOCX “Descontos Finais”
# -----------------------------------------------------------------
def gerar_docx_finais(df: pd.DataFrame, titulo: str) -> bytes:
    """
    Gera DOCX com as mesmas configurações do PDF "Descontos Finais":
    - Orientação paisagem;
    - Título principal centralizado;
    - Tabela com colunas: COD(30mm), DESCRIÇÃO(120mm), DESCONTOS(40mm), DATA(30mm);
    - Linhas especiais (A/B/Indébito) em vermelho, tamanho 11, negrito;
    - Converte valores de DESCONTOS para padrão PT-BR.
    """
    import re

    # Faz ajuste no título (eventual troca de vírgula por ponto, se quiser)
    # Aqui mantemos a substituição para casos como "14,642-0" -> "14.642-0"
    titulo_ajust = re.sub(r"(\d+),(\d{3}-\d\s*[A-Za-z])", r"\1.\2", titulo)

    doc = Document()
    for section in doc.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height

    # Título
    titulo_heading = doc.add_heading(titulo_ajust, level=1)
    titulo_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if df.empty:
        p = doc.add_paragraph("DataFrame vazio - nenhum dado para exibir.")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    col_names = ["COD", "DESCRIÇÃO", "DESCONTOS", "DATA"]
    widths_map = {
        "COD": 30,
        "DESCRIÇÃO": 120,
        "DESCONTOS": 40,
        "DATA": 30
    }

    table = doc.add_table(rows=1, cols=len(col_names))
    table.style = 'Table Grid'

    # Cabeçalho
    hdr_cells = table.rows[0].cells
    for i, col_h in enumerate(col_names):
        hdr_cells[i].text = col_h
        for parag in hdr_cells[i].paragraphs:
            for run in parag.runs:
                run.font.bold = True
                run.font.name = "Arial"
                run.font.size = Pt(10)

    # Linhas
    for _, rowv in df.iterrows():
        row_cells = table.add_row().cells
        desc_val = rowv.get("DESCRIÇÃO", "")
        is_especial = desc_val in [
            "A = Valor Total (R$)",
            "B = Valor Recebido - Autor (a)",
            "Indébito (A-B)",
            "Indébito em dobro (R$)"
        ]

        for j, col_h in enumerate(col_names):
            val = str(rowv.get(col_h, ""))
            if col_h == "DESCONTOS" and val.strip():
                val = formatar_valor_brl(val)

            para = row_cells[j].paragraphs[0]
            run = para.add_run(val)
            run.font.name = "Arial"

            if is_especial:
                run.font.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(255, 0, 0)
            else:
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0, 0, 0)

            if col_h == "DESCRIÇÃO":
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Ajusta larguras (mm -> inches)
    for i, col_h in enumerate(col_names):
        width_inches = widths_map[col_h] / 25.4
        table.columns[i].width = Inches(width_inches)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def ajustar_valores_docx(file_input_bytes: bytes) -> bytes:
    """
    Exemplo de função para varrer o DOCX e ajustar valores, se necessário.
    Aqui não está sendo usada, mas mantida como referência.
    """
    import re
    from docx import Document
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_in:
        tmp_in.write(file_input_bytes)
        tmp_in.flush()
        in_path = tmp_in.name
    out_path = in_path.replace(".docx", "_corrigido.docx")

    doc = Document(in_path)
    pattern = re.compile(r'([\d,]+\.\d{2})')
    for para in doc.paragraphs:
        found = pattern.findall(para.text)
        for val_us in found:
            try:
                base_float = float(val_us.replace(",", "").replace(".", "")) / 100
                val_br = f"{base_float:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
            except:
                val_br = val_us
            para.text = para.text.replace(val_us, val_br)

    doc.save(out_path)
    with open(out_path, "rb") as f:
        data_bytes = f.read()
    os.remove(in_path)
    os.remove(out_path)
    return data_bytes


# -----------------------------------------------------------------
# CRUZAR DESCONTOS COM GLOSSÁRIO
# -----------------------------------------------------------------
def cruzar_descontos_com_rubricas(df_descontos, rubricas, threshold=85):
    """
    Faz fuzzy matching das colunas 'DESCRIÇÃO' em df_descontos com a lista de rubricas,
    mantendo apenas as que atingirem a similaridade mínima (threshold).
    """
    if df_descontos.empty or not rubricas:
        return pd.DataFrame()

    unique_desc = df_descontos["DESCRIÇÃO"].unique()
    mapping = {}
    for desc in unique_desc:
        result = process.extractOne(desc, rubricas, scorer=fuzz.ratio)
        mapping[desc] = (result is not None and result[1] >= threshold)

    mask = df_descontos["DESCRIÇÃO"].map(mapping)
    return df_descontos[mask]


# -----------------------------------------------------------------
# MAIN
# -----------------------------------------------------------------
def main():
    # Exibir logomarca
    logo_base64 = get_image_base64(LOGO_PATH)
    if logo_base64:
        st.markdown(
            f"""
            <div style="text-align:center; margin-bottom:10px;">
                <img src="data:image/png;base64,{logo_base64}" alt="Logomarca" style="width:300px;" />
            </div>
            """,
            unsafe_allow_html=True
        )

    # Título principal (fonte reduzida)
    st.markdown("<h3 style='text-align:center;'>Analisador de Contracheques - SEAD / AMAZONPREV</h3>",
                unsafe_allow_html=True)

    # Upload do PDF
    uploaded_pdf = st.file_uploader(
        "Clique no botão para enviar o arquivo PDF (Contracheques - SEAD / AMAZONPREV)",
        type="pdf"
    )

    # Se o usuário enviou um PDF, processa...
    if uploaded_pdf is not None:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
            tmp_file.write(uploaded_pdf.read())
            pdf_path = tmp_file.name

        try:
            # Extrair nome e matrícula
            nome_final, nit_final = extrair_nome_e_matricula(pdf_path)
            set_state_value("nome_extraido", nome_final)
            set_state_value("nit_extraido", nit_final)

            # Ler tabelas Camelot
            df_contra = processar_contracheques_camelot(pdf_path)
            set_state_value("df_contracheques", df_contra)

        finally:
            # Remove o arquivo temporário do PDF original
            os.unlink(pdf_path)

    # Recupera DataFrame e nome/nit do estado
    df_contra = get_state_value("df_contracheques")
    nome_user = get_state_value("nome_extraido") or "ND"
    nit_user = get_state_value("nit_extraido") or "ND"

    # Se DataFrame não é vazio, exibe sucesso e segue fluxo
    if df_contra is not None and not df_contra.empty:
        st.success("PDF processado com sucesso!")

        # Tabelas (SEAD / AMAZONPREV)
        st.subheader("Tabelas (SEAD / AMAZONPREV)")
        st.dataframe(df_contra)

        # Botão de Download PDF dessas Tabelas
        pdf_tab_bytes = salvar_em_pdf_camelot(
            df_contra,
            "Tabelas (SEAD / AMAZONPREV)",
            nome_user,
            nit_user
        )
        pdf_tab_filename = f"Contracheque_{nome_user}_{nit_user}.pdf"
        st.download_button(
            label="Baixar PDF Tabelas (SEAD / AMAZONPREV)",
            data=pdf_tab_bytes,
            file_name=pdf_tab_filename,
            mime="application/pdf"
        )

        # Lista de Rubricas
        st.markdown("## Lista de Rubricas")
        rubricas_list = carregar_glossario_rubricas()
        if rubricas_list:
            df_rubricas = pd.DataFrame({"Rubricas": rubricas_list})
            st.dataframe(df_rubricas, use_container_width=True)
        else:
            st.warning("Glossário vazio ou não encontrado.")

        # Filtrar Descontos no Glossário
        st.markdown("## Filtrar Descontos no Glossário")
        with st.form("form_filtro_gloss"):
            thresh = st.slider("Nível de Similaridade (0.1 a 1.0)", 0.1, 1.0, 0.85, 0.1)
            btn_gloss = st.form_submit_button("Filtrar com Rubricas")

        if btn_gloss:
            # Ajusta colunas para padronizar => "DESCRIÇÃO", "DESCONTOS"
            df_aux = df_contra.copy()
            df_aux.rename(columns={
                "Descrição": "DESCRIÇÃO",
                "TOTAL": "DESCONTOS"
            }, inplace=True)
            df_aux = df_aux[df_aux["DESCONTOS"].str.strip() != ""]
            set_state_value("df_descontos", df_aux)

            threshold_val = int(thresh * 100)
            df_desc_gloss = cruzar_descontos_com_rubricas(df_aux, rubricas_list, threshold_val)
            set_state_value("df_descontos_gloss", df_desc_gloss)
            set_state_value("df_descontos_gloss_sel", None)

        df_desc_gloss = get_state_value("df_descontos_gloss")
        if df_desc_gloss is not None and not df_desc_gloss.empty:
            st.markdown("### Descontos x Glossário")
            st.dataframe(df_desc_gloss, use_container_width=True)

            # Lista única de descontos
            st.markdown("## Lista única de descontos")
            df_sel = get_state_value("df_descontos_gloss_sel")
            if df_sel is None:
                df_sel = df_desc_gloss

            with st.form("form_inclusao_descontos"):
                valores_unicos = sorted(df_sel["DESCRIÇÃO"].unique())
                st.write("Marque os itens que deseja incluir:")
                selected_descr = []
                for i, val_item in enumerate(valores_unicos):
                    qtd = df_sel[df_sel["DESCRIÇÃO"] == val_item].shape[0]
                    label_str = f"{i + 1} - {val_item} (Qtd: {qtd})"
                    if st.checkbox(label_str, key=f"chk_{i}"):
                        selected_descr.append(val_item)
                btn_incluir = st.form_submit_button("Confirmar Inclusão (Descontos)")

            if btn_incluir:
                if selected_descr:
                    df_incluido = df_sel[df_sel["DESCRIÇÃO"].isin(selected_descr)].copy()
                    set_state_value("df_descontos_gloss_sel", df_incluido)
                    st.success("Descontos selecionados com sucesso!")
                    st.markdown("### Lista restantes após exclusões")
                    st.dataframe(df_incluido, use_container_width=True)
                else:
                    st.warning("Nenhuma descrição selecionada.")

            # Apresentar Rúbricas para Débitos (Descontos Finais)
            df_final_sel = get_state_value("df_descontos_gloss_sel")
            if df_final_sel is not None and not df_final_sel.empty:
                st.markdown("## Apresentar Rúbricas para Débitos (Descontos Finais)")

                # Input para B = Valor Recebido
                valor_receb_input = st.text_input("B = Valor Recebido - Autor (a)", "0")
                set_state_value("valor_recebido", valor_receb_input)

                with st.form("form_descontos_finais"):
                    btn_final = st.form_submit_button("Gerar Relatório Final de Descontos")

                if btn_final:
                    df_final = df_final_sel.copy()
                    # Ordena por DATA, se existir
                    if "DATA" in df_final.columns:
                        df_final = df_final.sort_values(by="DATA").reset_index(drop=True)

                    # Insere as linhas A/B/Indébito
                    df_final = inserir_totais_na_coluna(df_final, "DESCONTOS")

                    # Gera PDF Finais
                    final_title = f"Descontos Finais - {nome_user} - {nit_user}"
                    pdf_fin_bytes = gerar_pdf_finais(df_final, final_title)
                    pdf_fin_name = f"Contracheque_Descontos_Finais_{nome_user}_{nit_user}.pdf"
                    st.download_button(
                        label="Baixar PDF (Descontos Finais)",
                        data=pdf_fin_bytes,
                        file_name=pdf_fin_name,
                        mime="application/pdf"
                    )

                    # Gera DOCX Finais
                    docx_bytes = gerar_docx_finais(df_final, final_title)
                    docx_fin_name = pdf_fin_name.replace(".pdf", ".docx")
                    st.download_button(
                        label="Baixar DOCX (Descontos Finais)",
                        data=docx_bytes,
                        file_name=docx_fin_name,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

    else:
        # Se df_contra não existe mas o usuário não enviou PDF ou processou sem sucesso
        if df_contra is not None:
            st.warning("Não foi possível extrair dados do PDF ou o PDF está vazio.")


if __name__ == "__main__":
    main()
